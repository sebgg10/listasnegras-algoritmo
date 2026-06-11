"""
Genera los documentos SRS y SDD para el proyecto
Módulo de Verificación de Nombres — Entidades Financieras (CNBV/PLD-FT)
v2 — incluye flujo completo de importación y comparación de oficios CNBV
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))


# ─────────────────────────────────────────
#  Helpers de formato
# ─────────────────────────────────────────

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_table_row(table, cells):
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = str(val)
    return row


def shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def make_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        hdr.cells[i].paragraphs[0].runs[0].bold = True
    shade_row(hdr, "D9E1F2")
    for row in rows:
        add_table_row(tbl, row)
    return tbl


# ═══════════════════════════════════════════════════════════════════
#  SRS
# ═══════════════════════════════════════════════════════════════════

def build_srs():
    doc = Document()

    t = doc.add_heading("Software Requirements Specification (SRS)", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Módulo de Verificación de Nombres")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)
    doc.add_paragraph("Empresa/Dependencia: Entidades Financieras (CNBV / PLD-FT)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Modelo de negocio: B2A").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tabla de versiones
    add_heading(doc, "Tabla de Control de Versiones", 1)
    make_table(doc,
        ["Versión", "Fecha", "Autor(es)", "Descripción del Cambio"],
        [
            ["1.0", "2026-06-11", "Equipo de Desarrollo", "Versión inicial: búsqueda manual difusa"],
            ["2.0", "2026-06-11", "Equipo de Desarrollo", "Incorpora importación de oficios CNBV (XML) y flujo de comparación automatizada"],
        ]
    )

    # ── 1. Introducción ──────────────────────────────────────────────
    add_heading(doc, "1. Introducción", 1)

    add_heading(doc, "1.1 Propósito del Documento", 2)
    doc.add_paragraph(
        "El presente documento define la visión general, el alcance funcional, los objetivos y las "
        "restricciones del prototipo search_name. Su función es alinear la propuesta técnica con la "
        "necesidad del proyecto: contar con una muestra funcional que permita buscar y comparar nombres "
        "provenientes de los oficios CNBV contra la base de datos interna de clientes de una entidad "
        "financiera, tolerando errores ortográficos, variaciones de captura, nombres incompletos y "
        "cambios en el orden de los componentes del nombre."
    )
    doc.add_paragraph(
        "Este prototipo responde directamente al mandato regulatorio de la Comisión Nacional Bancaria "
        "y de Valores (CNBV) a las entidades financieras: comparar los nombres contenidos en los "
        "oficios de Prevención de Lavado de Dinero y Financiamiento al Terrorismo (PLD-FT) contra su "
        "base de datos de clientes, a fin de detectar posibles coincidencias de forma oportuna."
    )

    add_heading(doc, "1.2 Alcance del Documento", 2)
    doc.add_paragraph("Este documento abarca:")
    for item in [
        "La visión general del prototipo y su contexto regulatorio.",
        "El problema que busca atender (comparación manual vs. algorítmica de nombres).",
        "Los actores involucrados y sus responsabilidades.",
        "Las características funcionales y no funcionales del sistema.",
        "Las restricciones actuales del proyecto.",
        "Los límites del alcance en esta etapa (prototipo, no integración final).",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.3 Definiciones y Acrónimos", 2)
    make_table(doc,
        ["Término / Acrónimo", "Definición"],
        [
            ("CNBV", "Comisión Nacional Bancaria y de Valores"),
            ("PLD", "Prevención de Lavado de Dinero"),
            ("FT", "Financiamiento al Terrorismo"),
            ("LD", "Lavado de Dinero"),
            ("UIF", "Unidad de Inteligencia Financiera — autoridad emisora de los oficios"),
            ("Oficio CNBV", "Documento oficial emitido por la CNBV durante días hábiles que notifica a las entidades financieras sobre personas bajo investigación por PLD/FT"),
            ("Lista de Personas Bloqueadas", "Lista oficial emitida por la UIF con personas físicas o morales sujetas a suspensión inmediata de operaciones"),
            ("B2A", "Business to Administration — modelo de negocio donde una empresa provee servicios a una entidad gubernamental o regulatoria"),
            ("RFC", "Registro Federal de Contribuyentes — clave fiscal mexicana"),
            ("Persona Física", "Individuo en la lista CNBV; identificado por nombre, apellido paterno y apellido materno"),
            ("Persona Moral", "Empresa o asociación en la lista CNBV; identificada por razón social en el campo Nombre"),
            ("Coincidencia difusa", "Técnica que encuentra similitudes entre cadenas de texto aunque no sean idénticas"),
            ("Umbral de coincidencia", "Porcentaje mínimo a partir del cual un resultado se considera relevante"),
            ("MVT", "Modelo-Vista-Template — patrón arquitectónico de Django"),
            ("SQLite", "Motor de base de datos relacional embebido utilizado en el prototipo"),
            ("rapidfuzz", "Biblioteca Python para comparación de cadenas con métricas de similitud textual"),
            ("jellyfish", "Biblioteca Python para comparación fonética y métricas de distancia de edición"),
            ("Metaphone", "Algoritmo fonético que codifica palabras según su pronunciación"),
            ("Levenshtein", "Métrica que mide el número mínimo de ediciones para transformar una cadena en otra"),
        ]
    )

    # ── 2. Posicionamiento ───────────────────────────────────────────
    add_heading(doc, "2. Posicionamiento del Producto", 1)

    add_heading(doc, "2.1 Oportunidad de Negocio", 2)
    doc.add_paragraph(
        "La CNBV emite diariamente oficios en formato XML donde informa a las entidades financieras "
        "sobre personas físicas y morales (nacionales y extranjeras) bajo investigación por posibles "
        "delitos de lavado de dinero o financiamiento al terrorismo. Cada entidad debe contrastar esos "
        "nombres contra su base de datos de clientes para detectar coincidencias y aplicar las medidas "
        "indicadas en el oficio (suspensión de operaciones, reporte de operación inusual, notificación al cliente)."
    )
    doc.add_paragraph("En la práctica, el proceso enfrenta los siguientes problemas:")
    for item in [
        "Errores de escritura o captura (Hernandez → Hernandes, Herdez, HDEZ).",
        "Acentos omitidos o inconsistentes (García → Garcia, Garzia).",
        "Cambio en el orden de nombre y apellidos (Lopez Juan → Juan Lopez).",
        "Nombres incompletos o abreviados.",
        "Nombres extranjeros con transliteraciones variables.",
        "Proceso actual 100% manual: un analista revisa los XMLs, extrae los nombres a mano y los compara visualmente.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.2 Declaración de Visión", 2)
    doc.add_paragraph(
        "Desarrollar un sistema web que permita importar automáticamente los oficios CNBV desde sus "
        "archivos XML, comparar las personas contenidas en cada oficio contra la base de datos interna "
        "de clientes mediante coincidencia difusa, y mostrar los resultados con porcentajes de similitud, "
        "de forma que la propuesta técnica sea demostrable, entendible y escalable hacia una solución "
        "institucional posterior."
    )

    add_heading(doc, "2.3 Objetivos del Negocio", 2)
    make_table(doc,
        ["ID", "Objetivo"],
        [
            ("OBJ-01", "Eliminar la comparación manual de nombres reemplazándola por un proceso automatizado de importación y matching."),
            ("OBJ-02", "Reducir el riesgo de omisión de coincidencias por errores ortográficos o variaciones en la captura."),
            ("OBJ-03", "Proveer un criterio objetivo y reproducible (porcentaje) que apoye la decisión del analista."),
            ("OBJ-04", "Automatizar la lectura de los archivos XML que la CNBV ya entrega a las entidades financieras."),
            ("OBJ-05", "Demostrar la viabilidad técnica del algoritmo ante la dirección y equipos de cumplimiento."),
            ("OBJ-06", "Sentar las bases documentadas para una integración futura con datos reales y sistemas legados."),
        ]
    )

    # ── 3. Stakeholders ──────────────────────────────────────────────
    add_heading(doc, "3. Stakeholders y Perfiles de Usuario", 1)

    add_heading(doc, "3.1 Stakeholders Clave", 2)
    make_table(doc,
        ["Stakeholder", "Rol / Interés", "Prioridad"],
        [
            ("Oficial de Cumplimiento", "Responsable del proceso PLD-FT. Define criterios de aceptación de coincidencias y recibe el reporte de alertas.", "Alta"),
            ("Analista de Cumplimiento", "Opera el sistema día a día: importa oficios, ejecuta comparaciones y escala alertas.", "Alta"),
            ("Área de Tecnología (TI)", "Despliega y mantiene el sistema; garantiza la integración eventual con datos reales.", "Media"),
            ("Dirección General / Auditoría", "Recibe reportes y valida que el proceso cumpla las exigencias regulatorias de la CNBV.", "Media"),
            ("CNBV / UIF (regulador)", "Emite los oficios XML con los nombres a verificar; no interactúa directamente con el sistema.", "Alta — fuente externa"),
        ]
    )

    add_heading(doc, "3.2 Clases de Usuarios", 2)
    make_table(doc,
        ["Tipo de Usuario", "Descripción", "Punto de Acceso"],
        [
            ("Analista de Cumplimiento", "Usuario primario. Importa XMLs, ejecuta comparaciones contra la base de clientes e interpreta porcentajes.", "Interfaz web: /oficios/ y /oficios/<folio>/comparar/"),
            ("Buscador manual", "Realiza búsquedas individuales de nombres para validaciones puntuales.", "Interfaz web: / (búsqueda difusa)"),
            ("Administrador del Sistema", "Gestiona la base de datos de clientes, carga/actualiza registros y accede al panel de administración Django.", "Panel /admin/ + comandos manage.py"),
        ]
    )

    # ── 4. Alcance ───────────────────────────────────────────────────
    add_heading(doc, "4. Alcance del Producto", 1)

    add_heading(doc, "4.1 Características Principales", 2)
    make_table(doc,
        ["ID", "Épica", "Descripción"],
        [
            ("EP-01", "Búsqueda difusa manual", "El analista ingresa un nombre y obtiene los registros más similares de la BD, tolerando errores ortográficos, fonéticos y de orden."),
            ("EP-02", "Porcentaje de coincidencia", "Cada resultado muestra un porcentaje calculado que combina métricas textuales y fonéticas."),
            ("EP-03", "Configuración del umbral", "El usuario ajusta el porcentaje mínimo de coincidencia y el número máximo de resultados."),
            ("EP-04", "Contexto del registro", "Los resultados incluyen RFC, origen del registro y folio de referencia."),
            ("EP-05", "Explicación del resultado", "El sistema indica el motivo principal de la coincidencia (ortografía, pronunciación, orden, etc.)."),
            ("EP-06", "Importación de oficios CNBV", "El sistema lee los archivos XML de la CNBV (estructurados bajo el namespace http://www.cnbv.gob.mx) y extrae el encabezado del oficio y los datos de cada persona (física o moral)."),
            ("EP-07", "Lista de oficios importados", "Página web que muestra todos los oficios importados con su folio, fecha, referencia, plazo (con indicador visual de urgencia) y conteo de personas."),
            ("EP-08", "Comparación oficio vs. base de clientes", "Para un oficio seleccionado, el sistema compara automáticamente todas sus personas contra la base completa de clientes y muestra los resultados agrupados por persona, con semáforo de alerta."),
            ("EP-09", "Base de datos local con datos ficticios", "Comando seed_names que carga registros de demostración para ilustrar el funcionamiento del algoritmo."),
            ("EP-10", "Panel de administración", "Administración básica de registros a través del panel Django Admin (NameRecord, Oficio, PersonaCNBV)."),
        ]
    )

    add_heading(doc, "4.2 Exclusiones (Out of Scope)", 2)
    for item in [
        "Integración con datos reales sensibles de clientes.",
        "Conexión operativa a bases de datos MySQL del sistema legado.",
        "Pruebas formales exhaustivas y certificación regulatoria.",
        "Despliegue productivo en infraestructura de la entidad financiera.",
        "Aplicación móvil o API REST pública.",
        "Generación automática de reportes para la CNBV.",
        "Descarga automática de oficios desde el portal CNBV (el prototipo asume que los XML ya están disponibles localmente).",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.3 Inclusiones del Alcance", 2)
    for item in [
        "Interfaz web de búsqueda difusa manual de nombres.",
        "Importación de oficios CNBV desde archivos XML locales (comando import_cnbv).",
        "Lista web de oficios importados con indicadores de urgencia.",
        "Comparación automatizada oficio-vs-clientes con resultados por porcentaje.",
        "Base de datos local SQLite con modelos estructurados (NameRecord, Oficio, PersonaCNBV).",
        "Panel de administración Django para gestión de registros.",
        "Documentación base (SRS + SDD) coherente con la implementación.",
    ]:
        add_bullet(doc, item)

    # ── 5. Restricciones ─────────────────────────────────────────────
    add_heading(doc, "5. Restricciones", 1)

    add_heading(doc, "5.1 Restricciones Técnicas", 2)
    for item in [
        "El prototipo se ejecuta sobre el framework Django 5.x (Python 3.11).",
        "La persistencia local se implementa con SQLite.",
        "Los datos de clientes utilizados son ficticios; no se procesa información real.",
        "La interfaz debe permanecer simple y funcional, sin diseño complejo.",
        "La lógica de matching debe estar desacoplada de la vista y del modelo de datos.",
        "Los XMLs de los oficios CNBV deben estar disponibles localmente en la carpeta de listasnegras.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.2 Restricciones de Proyecto", 2)
    for item in [
        "El prototipo debe quedar sólido y demostrable antes de cualquier integración con datos reales.",
        "La documentación formal debe respaldar el prototipo, no sustituirlo.",
        "No debe exponerse ni procesarse información sensible real en esta etapa.",
        "El desarrollo sigue el modelo B2A: el cliente final es una entidad pública regulada.",
    ]:
        add_bullet(doc, item)

    # ── 6. Supuestos ──────────────────────────────────────────────────
    add_heading(doc, "6. Supuestos", 1)
    for item in [
        "Los archivos XML de los oficios CNBV siguen la estructura del namespace http://www.cnbv.gob.mx con los elementos Expediente, SolicitudPartes y SolicitudEspecifica/PersonasSolicitud.",
        "Los nombres del problema real pueden llegar completos, incompletos o partidos entre campos.",
        "Una coincidencia alta no implica confirmación absoluta de identidad; es un insumo para el analista.",
        "El apellido materno puede estar ausente en la consulta o en el registro de la base de datos.",
        "La solución puede evolucionar hacia una integración mayor con datos y sistemas reales.",
    ]:
        add_bullet(doc, item)

    # ── 7. Riesgos ────────────────────────────────────────────────────
    add_heading(doc, "7. Riesgos Iniciales", 1)
    make_table(doc,
        ["ID", "Prob.", "Riesgo", "Impacto", "Mitigación"],
        [
            ("R-01", "Alta", "Falsos negativos en el matching", "Un nombre real en la lista CNBV no supera el umbral.", "Calibrar el umbral mínimo con casos reales antes del despliegue."),
            ("R-02", "Media", "Falsos positivos excesivos", "Demasiados resultados irrelevantes generan fatiga en el analista.", "Ajustar los pesos del algoritmo y el umbral por defecto."),
            ("R-03", "Media", "Cambio de estructura XML", "La CNBV puede modificar el formato de sus oficios.", "Centralizar el parser en import_cnbv.py con funciones configurables."),
            ("R-04", "Media", "Datos de demostración no representativos", "El seed no cubre variantes de nombres extranjeros.", "Ampliar el conjunto de datos ficticios con más variantes."),
            ("R-05", "Baja", "Escalabilidad", "Con grandes volúmenes de clientes el algoritmo puede ser lento.", "Implementar pre-filtrado por índice o trigrama antes del matching."),
        ]
    )

    # ── 8. Criterio de Éxito ──────────────────────────────────────────
    add_heading(doc, "8. Criterio de Éxito del Prototipo", 1)
    doc.add_paragraph("El prototipo se considera exitoso si:")
    for item in [
        "Permite buscar nombres manualmente con resultados razonables incluso con errores de escritura.",
        "Importa correctamente los archivos XML de los oficios CNBV sin intervención manual.",
        "Compara automáticamente todas las personas de un oficio contra la base de clientes.",
        "Muestra porcentajes de coincidencia útiles, ordenados de mayor a menor.",
        "Soporta errores comunes de captura (cambio de letras, acentos, orden de tokens).",
        "Puede explicarse y sustentarse en la documentación formal (SRS + SDD).",
    ]:
        add_bullet(doc, item)

    # ── 9. Requisitos Funcionales ─────────────────────────────────────
    add_heading(doc, "9. Requisitos Funcionales", 1)
    make_table(doc,
        ["ID", "Prioridad", "Nombre", "Descripción"],
        [
            ("RF-01", "Alta", "Interfaz de búsqueda manual", "El sistema provee una interfaz web con campo de texto para el nombre, control de umbral mínimo (0-100%) y control de límite de resultados (1-100)."),
            ("RF-02", "Alta", "Motor de coincidencia difusa", "El sistema compara el nombre ingresado contra todos los registros aplicando métricas textuales (ratio, token_sort_ratio, token_set_ratio, partial_ratio, Levenshtein, Jaro-Winkler) y fonéticas (Metaphone), combinadas en una puntuación ponderada 0-100."),
            ("RF-03", "Alta", "Porcentaje de coincidencia", "Cada resultado muestra el porcentaje calculado, redondeado a dos decimales."),
            ("RF-04", "Alta", "Filtrado por umbral", "El sistema excluye resultados con puntuación inferior al umbral mínimo configurado."),
            ("RF-05", "Alta", "Ordenamiento de resultados", "Los resultados se presentan de mayor a menor porcentaje de coincidencia."),
            ("RF-06", "Media", "Normalización de nombres", "Antes de comparar, el sistema normaliza: elimina acentos, convierte a minúsculas, elimina caracteres no alfanuméricos y colapsa espacios."),
            ("RF-07", "Media", "Contexto del registro", "Cada resultado incluye RFC, origen del registro y folio de referencia."),
            ("RF-08", "Media", "Explicación del resultado", "El sistema indica el motivo principal de la coincidencia (exacta, orden diferente, fonética, abreviación, tipografía o difusa combinada)."),
            ("RF-09", "Media", "Variantes de candidato", "El motor evalúa variantes del candidato (con/sin apellido materno) y retiene la de mayor similitud."),
            ("RF-10", "Alta", "Importación de oficios CNBV desde XML", "El sistema lee archivos XML con el namespace http://www.cnbv.gob.mx de forma recursiva desde una carpeta local, extrae los datos de cada oficio (folio, número, expediente, fecha, plazo, autoridad, referencia) y de cada PersonasSolicitud (PersonaId, Caracter, Persona, Paterno, Materno, Nombre, Rfc, Relacion, Domicilio, Complementarios). Los oficios ya importados no se duplican."),
            ("RF-11", "Alta", "Lista de oficios importados", "El sistema muestra una página web con todos los oficios importados: folio, fecha de publicación, referencia, área, plazo (con indicador visual: urgente ≤1 día / normal), número de personas y botón para iniciar la comparación."),
            ("RF-12", "Alta", "Comparación oficio vs. base de clientes", "Para un oficio seleccionado, el sistema compara automáticamente todas sus personas contra todos los NameRecord de la base de datos, agrupando los resultados por persona CNBV y mostrando un resumen (total personas, total clientes, total con coincidencias)."),
            ("RF-13", "Alta", "Indicadores visuales de alerta", "Las tarjetas de personas con al menos una coincidencia por encima del umbral se resaltan visualmente. Las filas de resultados usan color semáforo: rojo (≥80%), naranja (≥60%), amarillo (<60%)."),
            ("RF-14", "Media", "Umbral ajustable en comparación", "En la página de comparación el analista puede ajustar el umbral y recalcular sin salir de la vista."),
            ("RF-15", "Baja", "Carga de datos de demostración", "El comando seed_names carga 24 registros ficticios diseñados para ilustrar los diferentes escenarios de matching."),
            ("RF-16", "Baja", "Panel de administración", "El sistema expone /admin/ con gestión de NameRecord, Oficio y PersonaCNBV."),
        ]
    )

    # ── 10. Requisitos No Funcionales ─────────────────────────────────
    add_heading(doc, "10. Requisitos No Funcionales", 1)
    make_table(doc,
        ["ID", "Nombre", "Descripción"],
        [
            ("RNF-01", "Seguridad de datos", "El prototipo no debe requerir ni exponer datos reales sensibles de clientes. En producción se requiere autenticación y cifrado en tránsito."),
            ("RNF-02", "Mantenibilidad", "La lógica de matching (matching.py) y el parser XML (import_cnbv.py) deben estar desacoplados de vistas y modelos."),
            ("RNF-03", "Claridad de demostración", "La interfaz debe ser comprensible para un analista de cumplimiento sin formación técnica."),
            ("RNF-04", "Portabilidad", "El sistema debe ejecutarse localmente en cualquier entorno con Python 3.11 y Django 5.x, sin dependencias de infraestructura externa."),
            ("RNF-05", "Trazabilidad regulatoria", "En versiones productivas, cada consulta y comparación debe quedar registrada con fecha, usuario y resultado para auditoría por la CNBV."),
        ]
    )

    # ── 11. Reglas de Negocio ─────────────────────────────────────────
    add_heading(doc, "11. Reglas de Negocio del Prototipo", 1)
    make_table(doc,
        ["ID", "Regla"],
        [
            ("RN-01", "El porcentaje final se expresa entre 0 y 100 (inclusive)."),
            ("RN-02", "Los resultados se muestran solo si superan el umbral mínimo configurado."),
            ("RN-03", "El nombre buscable (search_full_name) se construye evitando duplicar tokens repetidos entre nombre, paterno y materno."),
            ("RN-04", "Una coincidencia exacta tras normalización se representa con el valor máximo: 100%."),
            ("RN-05", "Registros con los mismos tokens en diferente orden reciben al menos 97% de coincidencia."),
            ("RN-06", "Los datos de clientes utilizados en el prototipo deben ser ficticios; no se procesan datos reales."),
            ("RN-07", "El sistema no debe duplicar un oficio ya importado (unicidad por folio)."),
            ("RN-08", "Para personas morales (Persona=Moral), el nombre buscable se construye únicamente con el campo Nombre (razón social); Paterno y Materno son vacíos."),
            ("RN-09", "Para personas físicas (Persona=Fisica), el nombre buscable se construye combinando Nombre, Paterno y Materno."),
        ]
    )

    # ── 12. Criterio de Aceptación ────────────────────────────────────
    add_heading(doc, "12. Criterio de Aceptación del Prototipo", 1)
    doc.add_paragraph("El prototipo se considera aceptable si:")
    for item in [
        "Importa correctamente los 64 oficios XML de la carpeta listasnegras/ sin errores.",
        "Muestra la lista de oficios con folio, fecha, plazo e indicador visual de urgencia.",
        "Compara cada persona CNBV contra la base de clientes y muestra resultados con porcentaje.",
        "Detecta variantes conocidas: HERNANDEZ → HERNANDES, HERDEZ, HDEZ.",
        "Detecta nombres en orden distinto: 'JUAN PEREZ LOPEZ' vs. 'PEREZ LOPEZ JUAN'.",
        "Muestra contexto básico (RFC, origen, folio) de cada resultado.",
        "Opera correctamente con datos ficticios sin requerir datos reales.",
        "La propuesta técnica queda documentada (SRS + SDD) y es coherente con la implementación.",
        "La interfaz es comprensible para un analista de cumplimiento sin formación técnica.",
    ]:
        add_bullet(doc, item)

    path = os.path.join(BASE, "Software Requirements Specification (SRS).docx")
    doc.save(path)
    print(f"SRS guardado: {path}")


# ═══════════════════════════════════════════════════════════════════
#  SDD
# ═══════════════════════════════════════════════════════════════════

def build_sdd():
    doc = Document()

    t = doc.add_heading("Software Design Document (SDD)", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Módulo de Verificación de Nombres — search_name")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)
    doc.add_paragraph("Empresa/Dependencia: Entidades Financieras (CNBV / PLD-FT)").alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Tabla de Control de Versiones", 1)
    make_table(doc,
        ["Versión", "Fecha", "Autor(es)", "Descripción del Cambio"],
        [
            ["1.0", "2026-06-11", "Equipo de Desarrollo", "Versión inicial: motor de matching difuso, modelo NameRecord"],
            ["2.0", "2026-06-11", "Equipo de Desarrollo", "Incorpora modelos Oficio y PersonaCNBV, comando import_cnbv, vistas lista_oficios y comparar_oficio"],
        ]
    )

    # ── 1. Propósito ──────────────────────────────────────────────────
    add_heading(doc, "1. Propósito del Diseño", 1)
    doc.add_paragraph(
        "Objetivo: Definir el diseño técnico del prototipo search_name, describiendo su estructura "
        "interna, sus componentes, sus decisiones de implementación y la forma en que el sistema "
        "resuelve tanto la búsqueda difusa manual de nombres como la comparación automatizada de "
        "oficios CNBV contra la base de clientes en el contexto de cumplimiento PLD-FT."
    )
    doc.add_paragraph(
        "Alcance: Este diseño cubre el prototipo funcional construido sobre Django 5.x con tres modelos "
        "(NameRecord, Oficio, PersonaCNBV), un motor de matching textual y fonético, un parser de "
        "XMLs CNBV, tres vistas web y un conjunto de datos ficticios para demostración. No cubre la "
        "integración final con datos reales ni la conexión operativa al sistema legado."
    )

    # ── 2. Arquitectura ───────────────────────────────────────────────
    add_heading(doc, "2. Arquitectura General", 1)

    add_heading(doc, "2.1 Estilo Arquitectónico", 2)
    doc.add_paragraph(
        "El sistema adopta el patrón nativo de Django: Modelo-Vista-Template (MVT). "
        "La separación de responsabilidades es la siguiente:"
    )
    for item in [
        "Modelo (models.py): define la estructura de datos y la persistencia en SQLite.",
        "Vista (views.py): coordina la lógica; recibe la petición HTTP, invoca el motor de matching o el motor de comparación, y envía los resultados al template.",
        "Template (*.html): presenta los resultados al usuario en HTML.",
        "Módulo de matching (matching.py): lógica de comparación difusa, completamente desacoplada.",
        "Comando de importación (import_cnbv.py): parser XML y carga de oficios, desacoplado de la vista.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.2 Diagrama de Contexto (texto)", 2)
    doc.add_paragraph(
        "El analista de cumplimiento accede desde un navegador web al sistema Django. "
        "Puede realizar dos flujos principales:\n"
        "(A) Búsqueda manual: ingresa un nombre en / y el sistema compara contra NameRecord.\n"
        "(B) Flujo CNBV: el administrador ejecuta 'python manage.py import_cnbv <carpeta>' para "
        "cargar los XMLs; luego el analista accede a /oficios/, selecciona un oficio y navega a "
        "/oficios/<folio>/comparar/ para ver los resultados automáticos. "
        "No existe comunicación con sistemas externos en el prototipo."
    )

    add_heading(doc, "2.3 Diagrama de Contenedores (texto)", 2)
    make_table(doc,
        ["Contenedor", "Responsabilidad", "Tecnología"],
        [
            ("Navegador web", "Cliente HTTP; no requiere JavaScript.", "Cualquier navegador moderno"),
            ("Aplicación Django (search_name)", "Servidor web de desarrollo; expone /, /oficios/, /oficios/<folio>/comparar/ y /admin/.", "Python 3.11 + Django 5.2.5"),
            ("Base de datos SQLite", "Archivo db.sqlite3 local; tablas coincidencias_namerecord, coincidencias_oficio, coincidencias_personacnbv.", "SQLite 3"),
            ("Carpeta listasnegras/", "Repositorio local de archivos XML de oficios CNBV. Solo se lee durante el import.", "Sistema de archivos"),
        ]
    )

    add_heading(doc, "2.4 Rutas URL", 2)
    make_table(doc,
        ["URL", "Vista / Módulo", "Nombre", "Descripción"],
        [
            ("/", "coincidencias.views.search_names", "search_names", "Búsqueda difusa manual de nombres (GET)."),
            ("/oficios/", "coincidencias.views.lista_oficios", "lista_oficios", "Lista de oficios CNBV importados."),
            ("/oficios/<folio>/comparar/", "coincidencias.views.comparar_oficio", "comparar_oficio", "Comparación de personas del oficio vs. base de clientes (GET, parámetro min_score)."),
            ("/admin/", "django.contrib.admin", "admin:index", "Panel de administración Django."),
        ]
    )

    # ── 3. Componentes Principales ────────────────────────────────────
    add_heading(doc, "3. Componentes Principales", 1)

    add_heading(doc, "3.1 models.py — Tres modelos", 2)
    doc.add_paragraph("El módulo define tres modelos persistidos en SQLite.")

    doc.add_paragraph("Modelo NameRecord — base de datos de clientes de la entidad financiera:").runs[0].bold = True
    make_table(doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ("nombre", "CharField(120)", "Nombre(s) de pila."),
            ("paterno", "CharField(80)", "Apellido paterno."),
            ("materno", "CharField(80)", "Apellido materno (puede vacío)."),
            ("rfc", "CharField(20)", "RFC (opcional)."),
            ("origen", "CharField(40)", "Fuente del registro."),
            ("folio_ref", "CharField(50)", "Referencia del oficio asociado."),
            ("search_full_name", "CharField(180) INDEX", "Nombre buscable derivado (auto)."),
            ("normalized_name", "CharField(180) INDEX", "Nombre normalizado derivado (auto)."),
            ("created_at / updated_at", "DateTimeField", "Timestamps de creación y modificación."),
        ]
    )

    doc.add_paragraph("Modelo Oficio — encabezado de cada oficio CNBV:").runs[0].bold = True
    make_table(doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ("folio", "IntegerField UNIQUE INDEX", "Número de folio del oficio (clave de unicidad)."),
            ("numero_oficio", "CharField(80)", "Número oficial (ej. 214-4/27921410/2026)."),
            ("numero_expediente", "CharField(80)", "Número de expediente interno CNBV."),
            ("solicitud_siara", "CharField(60)", "Clave SIARA (ej. UIFB/2026/001147)."),
            ("anio", "IntegerField", "Año del oficio."),
            ("area_clave / area_descripcion", "CharField", "Clave y descripción del área CNBV."),
            ("fecha_publicacion", "DateField", "Fecha en que se publicó el oficio."),
            ("dias_plazo", "IntegerField", "Días de plazo para responder (usualmente 1)."),
            ("autoridad_nombre", "CharField(220)", "Nombre de la autoridad emisora (UIF)."),
            ("referencia / referencia1", "CharField(220)", "Tipo de solicitud y clave de referencia."),
            ("tiene_aseguramiento", "BooleanField", "Si el oficio incluye aseguramiento."),
            ("archivo_xml", "CharField(220)", "Nombre del archivo XML de origen."),
            ("importado_en", "DateTimeField(auto_now_add)", "Timestamp de importación."),
        ]
    )

    doc.add_paragraph("Modelo PersonaCNBV — cada persona listada en un oficio:").runs[0].bold = True
    make_table(doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ("oficio", "ForeignKey(Oficio)", "Oficio al que pertenece esta persona."),
            ("persona_id", "IntegerField", "ID de persona dentro del oficio."),
            ("caracter", "CharField(60)", "Carácter (ej. Investigado)."),
            ("tipo_persona", "CharField(20)", "Tipo: 'Fisica' o 'Moral'."),
            ("nombre", "CharField(220)", "Nombre(s) de pila o razón social (moral)."),
            ("paterno", "CharField(80)", "Apellido paterno (vacío para morales)."),
            ("materno", "CharField(80)", "Apellido materno (vacío para morales)."),
            ("rfc", "CharField(20)", "RFC (puede vacío en SolicitudPartes, completo en PersonasSolicitud)."),
            ("relacion", "CharField(120)", "Relación con el investigado principal."),
            ("domicilio", "CharField(320)", "Domicilio registrado."),
            ("complementarios", "CharField(220)", "Datos complementarios (ej. ACUERDO: 173/2026)."),
            ("search_full_name", "CharField(300) INDEX", "Nombre buscable derivado (auto)."),
            ("normalized_name", "CharField(300) INDEX", "Nombre normalizado derivado (auto)."),
        ]
    )
    doc.add_paragraph(
        "Los tres modelos implementan prepare_search_fields() invocado automáticamente en save(), "
        "garantizando que search_full_name y normalized_name siempre estén actualizados."
    )

    add_heading(doc, "3.2 matching.py — Motor de Coincidencia", 2)
    doc.add_paragraph("Módulo central. Completamente desacoplado de Django; puede probarse en forma aislada.")
    make_table(doc,
        ["Función", "Responsabilidad"],
        [
            ("clean_name_piece(value)", "Elimina espacios múltiples y caracteres de control."),
            ("_dedupe_name_parts(parts)", "Elimina partes duplicadas de una lista comparando versiones normalizadas."),
            ("build_search_name(nombre, paterno, materno)", "Construye el nombre buscable evitando repetir tokens ya presentes en otro campo."),
            ("normalize_name(name)", "Convierte a minúsculas, elimina acentos (unidecode), elimina caracteres no alfanuméricos, colapsa espacios."),
            ("tokens(name)", "Devuelve la lista de tokens del nombre normalizado."),
            ("phonetic_token(token)", "Aplica Metaphone a un token; devuelve su código fonético."),
            ("_token_overlap_score(tokens_a, tokens_b)", "Porcentaje de tokens comunes entre dos listas (0-100)."),
            ("_phonetic_scores(query, candidate)", "Calcula phonetic_token_sort, phonetic_overlap y su combinación ponderada (70/30)."),
            ("score_name(query, candidate)", "Calcula todas las métricas y devuelve un dict con 'overall' y métricas individuales."),
            ("explain_score(metrics)", "Devuelve texto explicando el motivo principal de la coincidencia."),
            ("find_name_matches(query, candidates, min_score, limit)", "Orquesta la búsqueda: itera candidatos, calcula score, filtra por umbral, ordena y limita."),
        ]
    )

    add_heading(doc, "3.3 import_cnbv.py — Comando de Importación XML", 2)
    doc.add_paragraph(
        "Comando Django (python manage.py import_cnbv <carpeta>) que implementa el flujo completo "
        "de importación de oficios CNBV."
    )
    for item in [
        "Acepta la ruta de la carpeta raíz como argumento posicional.",
        "La bandera --reset elimina todos los oficios existentes antes de importar.",
        "Escanea recursivamente buscando archivos que coincidan con el patrón 4-*.xml.",
        "Parsea cada XML con xml.etree.ElementTree usando el namespace http://www.cnbv.gob.mx.",
        "Función auxiliar _text(element, tag_name): extrae el texto de un subelemento con fallback.",
        "Crea un registro Oficio con todos los campos del encabezado.",
        "Itera SolicitudEspecifica/PersonasSolicitud y crea un PersonaCNBV por cada elemento.",
        "Omite folios ya existentes en la BD (unicidad por folio).",
        "Reporta en consola: folio, fecha, cantidad de personas y resumen final (importados/omitidos/errores).",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.4 views.py — Tres Vistas", 2)

    doc.add_paragraph("search_names (GET /)").runs[0].bold = True
    for item in [
        "Instancia NameSearchForm con los parámetros GET.",
        "Obtiene query, min_score y limit mediante cleaned_or_default().",
        "Consulta NameRecord con .values() y llama a find_name_matches().",
        "Renderiza home.html con resultados y total de registros.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph("lista_oficios (GET /oficios/)").runs[0].bold = True
    for item in [
        "Consulta todos los Oficio con prefetch_related('personas') ordenados por fecha y folio.",
        "Renderiza oficios.html con la lista y el total de oficios.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph("comparar_oficio (GET /oficios/<folio>/comparar/)").runs[0].bold = True
    for item in [
        "Obtiene el Oficio por folio (404 si no existe).",
        "Lee el parámetro min_score del GET (default 45, rango 0-100).",
        "Consulta todos los NameRecord con .values().",
        "Por cada PersonaCNBV del oficio: llama a find_name_matches(persona.search_full_name, clientes, min_score, limit=10).",
        "Construye una lista de resultados con persona, matches y bandera tiene_coincidencias.",
        "Cuenta total_alertas (personas con al menos un match).",
        "Renderiza comparar.html con el resumen y los resultados por persona.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.5 forms.py — NameSearchForm", 2)
    doc.add_paragraph("Formulario Django (GET) con tres campos:")
    for item in [
        "query (CharField, max_length=180, opcional): nombre a buscar.",
        "min_score (IntegerField, 0-100, default=45): umbral mínimo de coincidencia.",
        "limit (IntegerField, 1-100, default=20): número máximo de resultados.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.6 Templates", 2)
    make_table(doc,
        ["Template", "Vista que lo usa", "Contenido principal"],
        [
            ("home.html", "search_names", "Formulario de búsqueda + tabla de resultados con porcentaje, motivo y métricas. Enlace a /oficios/."),
            ("oficios.html", "lista_oficios", "Tabla de oficios con folio, fecha, referencia, plazo (indicador color), conteo de personas y botón Comparar. Instrucción de carga si no hay oficios."),
            ("comparar.html", "comparar_oficio", "Encabezado del oficio + resumen (total personas, total clientes, total alertas). Por cada PersonaCNBV: tipo badge, nombre, RFC, carácter; tabla de clientes coincidentes con semáforo de color (rojo ≥80%, naranja ≥60%, amarillo <60%). Control de umbral ajustable."),
        ]
    )

    add_heading(doc, "3.7 admin.py — Tres AdminModels", 2)
    make_table(doc,
        ["AdminModel", "list_display", "search_fields", "list_filter"],
        [
            ("NameRecordAdmin", "full_name, normalized_name, created_at", "full_name, normalized_name", "created_at"),
            ("OficioAdmin", "folio, numero_oficio, fecha_publicacion, referencia, dias_plazo, importado_en", "folio, numero_oficio, numero_expediente, referencia", "fecha_publicacion, area_descripcion, tiene_aseguramiento"),
            ("PersonaCNBVAdmin", "search_full_name, tipo_persona, rfc, caracter, oficio", "search_full_name, normalized_name, rfc", "tipo_persona, caracter, oficio__fecha_publicacion"),
        ]
    )

    add_heading(doc, "3.8 seed_names.py — Datos de Demostración", 2)
    doc.add_paragraph(
        "Comando (python manage.py seed_names) que inserta 24 registros ficticios en NameRecord "
        "para demostrar los escenarios de matching. Incluye: pares con errores tipográficos "
        "(Juan/Johan, Perez/Peres), variantes fonéticas (Cristian/Christian, Hernandez/Hernandes), "
        "abreviaciones (Luis Eduardo/Luis Eduard) y casos con nombre en un solo campo vs. separado. "
        "La bandera --reset elimina todos los registros antes de insertar."
    )

    # ── 4. Base de Datos ──────────────────────────────────────────────
    add_heading(doc, "4. Diseño de Base de Datos", 1)

    add_heading(doc, "4.1 Tablas", 2)
    doc.add_paragraph("coincidencias_namerecord (base de clientes):").runs[0].bold = True
    make_table(doc,
        ["Columna", "Tipo / Restricción", "Descripción"],
        [
            ("id", "INTEGER PK AUTOINCREMENT", "Identificador único."),
            ("nombre", "VARCHAR(120)", "Nombre(s) de pila."),
            ("paterno", "VARCHAR(80)", "Apellido paterno."),
            ("materno", "VARCHAR(80)", "Apellido materno (puede vacío)."),
            ("rfc", "VARCHAR(20)", "RFC."),
            ("origen", "VARCHAR(40)", "Fuente del registro."),
            ("folio_ref", "VARCHAR(50)", "Referencia del oficio."),
            ("search_full_name", "VARCHAR(180) INDEX", "Nombre buscable (campo derivado)."),
            ("normalized_name", "VARCHAR(180) INDEX", "Nombre normalizado (campo derivado)."),
            ("created_at / updated_at", "DATETIME", "Timestamps."),
        ]
    )

    doc.add_paragraph("coincidencias_oficio (oficios CNBV):").runs[0].bold = True
    make_table(doc,
        ["Columna", "Tipo / Restricción", "Descripción"],
        [
            ("id", "INTEGER PK AUTOINCREMENT", "Identificador único."),
            ("folio", "INTEGER UNIQUE INDEX", "Folio del oficio (clave de negocio)."),
            ("numero_oficio", "VARCHAR(80)", "Número oficial."),
            ("numero_expediente", "VARCHAR(80)", "Número de expediente."),
            ("solicitud_siara", "VARCHAR(60)", "Clave SIARA."),
            ("anio", "INTEGER", "Año del oficio."),
            ("area_clave / area_descripcion", "VARCHAR", "Área CNBV."),
            ("fecha_publicacion", "DATE", "Fecha del oficio."),
            ("dias_plazo", "INTEGER", "Días de plazo de respuesta."),
            ("autoridad_nombre", "VARCHAR(220)", "Nombre de la autoridad."),
            ("referencia / referencia1", "VARCHAR(220)", "Tipo y clave de referencia."),
            ("tiene_aseguramiento", "BOOLEAN", "Si hay aseguramiento."),
            ("archivo_xml", "VARCHAR(220)", "Nombre del XML de origen."),
            ("importado_en", "DATETIME", "Timestamp de importación."),
        ]
    )

    doc.add_paragraph("coincidencias_personacnbv (personas por oficio):").runs[0].bold = True
    make_table(doc,
        ["Columna", "Tipo / Restricción", "Descripción"],
        [
            ("id", "INTEGER PK AUTOINCREMENT", "Identificador único."),
            ("oficio_id", "INTEGER FK → coincidencias_oficio", "Oficio al que pertenece."),
            ("persona_id", "INTEGER", "ID de persona dentro del oficio."),
            ("caracter", "VARCHAR(60)", "Carácter (Investigado, etc.)."),
            ("tipo_persona", "VARCHAR(20)", "Fisica o Moral."),
            ("nombre / paterno / materno", "VARCHAR", "Datos de identidad."),
            ("rfc", "VARCHAR(20)", "RFC."),
            ("relacion / domicilio / complementarios", "VARCHAR", "Datos adicionales del XML."),
            ("search_full_name", "VARCHAR(300) INDEX", "Nombre buscable (auto)."),
            ("normalized_name", "VARCHAR(300) INDEX", "Nombre normalizado (auto)."),
        ]
    )

    add_heading(doc, "4.2 Razón del Diseño con Modelos Separados", 2)
    doc.add_paragraph(
        "Se decidió separar Oficio y PersonaCNBV en lugar de una tabla única porque:"
    )
    for item in [
        "Un oficio puede contener múltiples personas (relación 1:N).",
        "Los datos del encabezado del oficio son estables y no deben repetirse por persona.",
        "La clave de unicidad del oficio (folio) es independiente de las personas.",
        "Este diseño replica la estructura lógica del XML original de la CNBV.",
        "El modelo PersonaCNBV reutiliza el mismo mecanismo de campos derivados (search_full_name, normalized_name) que NameRecord, garantizando comparación homogénea.",
    ]:
        add_bullet(doc, item)

    # ── 5. Flujo de la Aplicación ─────────────────────────────────────
    add_heading(doc, "5. Flujo de la Aplicación", 1)

    doc.add_paragraph("Flujo A — Búsqueda difusa manual:").runs[0].bold = True
    make_table(doc,
        ["Paso", "Descripción"],
        [
            ("1", "El analista ingresa un nombre en / y ajusta umbral y límite."),
            ("2", "HTTP GET → / con parámetros query, min_score, limit."),
            ("3", "search_names() valida el formulario y consulta NameRecord.objects.values()."),
            ("4", "find_name_matches() calcula el score de cada candidato."),
            ("5", "Se filtran por umbral, se ordenan y se limitan."),
            ("6", "home.html renderiza la tabla de resultados."),
        ]
    )

    doc.add_paragraph("Flujo B — Comparación de oficio CNBV:").runs[0].bold = True
    make_table(doc,
        ["Paso", "Actor", "Descripción"],
        [
            ("1", "Admin", "Ejecuta: python manage.py import_cnbv <carpeta_listasnegras>"),
            ("2", "import_cnbv", "Escanea 4-*.xml recursivamente, parsea namespace CNBV, crea Oficio + PersonaCNBV por cada archivo nuevo."),
            ("3", "Analista", "Navega a /oficios/ — ve la lista de oficios con folios, fechas y conteo de personas."),
            ("4", "Analista", "Hace clic en 'Comparar' para el oficio de interés → /oficios/<folio>/comparar/"),
            ("5", "comparar_oficio()", "Obtiene todas las PersonaCNBV del oficio y todos los NameRecord como candidatos."),
            ("6", "comparar_oficio()", "Por cada PersonaCNBV: llama a find_name_matches(persona.search_full_name, clientes, min_score, 10)."),
            ("7", "comparar_oficio()", "Construye lista de resultados con tiene_coincidencias y cuenta total_alertas."),
            ("8", "comparar.html", "Muestra resumen (personas, clientes, alertas), tarjetas por persona con semáforo de color y tabla de clientes coincidentes."),
            ("9", "Analista", "Ajusta el umbral en la misma página y recalcula si es necesario."),
        ]
    )

    # ── 6. Diseño del Algoritmo ───────────────────────────────────────
    add_heading(doc, "6. Diseño del Algoritmo de Matching", 1)

    add_heading(doc, "6.1 Preparación del Nombre (build_search_name)", 2)
    for item in [
        "Se limpian espacios múltiples de cada pieza (clean_name_piece).",
        "Se unen nombre, paterno y materno cuando están disponibles.",
        "Si los tokens de paterno o materno ya están en nombre (campo completo), se omiten para no duplicar.",
        "Se eliminan partes duplicadas comparando sus versiones normalizadas (_dedupe_name_parts).",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6.2 Normalización (normalize_name)", 2)
    for item in [
        "unidecode(): convierte Unicode a ASCII (elimina acentos, ñ → n).",
        "lower(): convierte a minúsculas.",
        "NON_ALNUM_RE.sub(): elimina todo lo que no sea letra, dígito o espacio.",
        "SPACE_RE.sub(): colapsa múltiples espacios en uno.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6.3 Métricas Textuales (rapidfuzz) y sus Pesos", 2)
    make_table(doc,
        ["Métrica", "Peso", "Justificación"],
        [
            ("token_sort_ratio", "0.23", "Ordena tokens antes de comparar. Mayor peso: resuelve el problema de orden de apellidos, el caso más frecuente en PLD-FT."),
            ("ratio", "0.15", "Similitud carácter a carácter. Sensible al orden y la posición."),
            ("token_set_ratio", "0.15", "Usa la intersección de tokens. Robusto ante nombres con partes extra o faltantes."),
            ("partial_ratio", "0.15", "Busca la subsecuencia más similar. Útil para nombres incompletos o abreviados."),
            ("Levenshtein normalizado", "0.12", "Número de ediciones mínimas normalizado. Complementa ratio."),
            ("Jaro-Winkler", "0.10", "Prioriza coincidencias al inicio de cadena. Útil para nombres propios con prefijo estable."),
            ("Fonética (Metaphone)", "0.10", "token_sort_ratio de códigos fonéticos (0.70) + overlap (0.30). Detecta variantes como HERNANDEZ/HERNANDES."),
        ]
    )
    doc.add_paragraph("Suma total de pesos: 1.00")

    add_heading(doc, "6.4 Ajustes del Ranking", 2)
    for item in [
        "Penalización por longitud: si la proporción de longitudes es <0.45 y partial_ratio >90, se aplica penalización para evitar que un string muy corto infle el score por ser subcadena.",
        "Coincidencia exacta tras normalización → score = 100.0",
        "Mismos tokens en diferente orden → score = max(score, 97.0)",
    ]:
        add_bullet(doc, item)

    # ── 7. Decisiones Clave ───────────────────────────────────────────
    add_heading(doc, "7. Decisiones Clave de Diseño", 1)
    make_table(doc,
        ["ID", "Decisión", "Justificación"],
        [
            ("DC-01", "Separar matching.py de models.py y views.py", "Permite probar el motor de forma aislada. Una versión futura puede reemplazar rapidfuzz por un modelo de ML sin tocar la vista."),
            ("DC-02", "Campos separados en NameRecord y PersonaCNBV (nombre/paterno/materno)", "Los oficios CNBV llegan con nombre partido entre campos. Separar permite construir variantes de candidato para maximizar la detección."),
            ("DC-03", "Múltiples métricas ponderadas", "Ninguna métrica aislada cubre todos los casos. La combinación hace el sistema más robusto frente a errores de captura, orden y fonética."),
            ("DC-04", "token_sort_ratio con mayor peso (0.23)", "El caso más frecuente en PLD-FT es el cambio de orden (apellido antes que nombre). Darle el mayor peso favorece la detección del escenario de mayor riesgo regulatorio."),
            ("DC-05", "Variantes de candidato (_candidate_variants)", "Un cliente puede estar registrado con o sin apellido materno. Evaluar ambas variantes garantiza que no se pierda una coincidencia real."),
            ("DC-06", "Parseo XML con ElementTree + namespace explícito", "La CNBV usa el namespace http://www.cnbv.gob.mx. Usar ElementTree nativo evita dependencias extra; el helper _tag() centraliza el prefijo de namespace."),
            ("DC-07", "Modelo Oficio separado de PersonaCNBV", "Replica la estructura 1:N del XML original (un oficio, múltiples personas). Facilita mostrar el encabezado del oficio independientemente de las personas."),
            ("DC-08", "SQLite para el prototipo", "Sin dependencias externas; facilita la portabilidad y la demostración. En producción se migrará a PostgreSQL o MySQL."),
            ("DC-09", "Umbral por defecto de 45%", "Valor empírico que equilibra falsos positivos y negativos en los datos de demostración. Debe recalibrarse con datos reales."),
        ]
    )

    # ── 8. Supuestos ──────────────────────────────────────────────────
    add_heading(doc, "8. Supuestos", 1)
    for item in [
        "Los archivos XML de los oficios CNBV siguen el namespace http://www.cnbv.gob.mx con la estructura Expediente > SolicitudEspecifica > PersonasSolicitud.",
        "Los nombres pueden llegar completos o fragmentados entre los campos del registro.",
        "El apellido materno puede faltar en la consulta del analista o en el registro de la BD.",
        "Una coincidencia alta no confirma por sí sola identidad; es un insumo para la decisión del analista.",
        "Los pesos del algoritmo son preliminares; requieren validación con datos reales.",
    ]:
        add_bullet(doc, item)

    # ── 9. Limitaciones ───────────────────────────────────────────────
    add_heading(doc, "9. Limitaciones Conocidas", 1)
    for item in [
        "Sin descarga automática: el sistema asume que los XML ya están disponibles localmente. La descarga desde el portal CNBV no está implementada.",
        "Sin conexión con datos reales de clientes.",
        "Sin autenticación: en el prototipo cualquiera con acceso puede consultar. En producción se requiere gestión de sesiones.",
        "Sin registro de auditoría: las consultas y comparaciones no se persisten. Necesario en producción para trazabilidad regulatoria.",
        "Escalabilidad: el algoritmo itera todos los registros en memoria. Para bases grandes se requiere una etapa previa de pre-filtrado (índice de trigramas o búsqueda vectorial).",
        "Los pesos del algoritmo pueden refinarse con muestras reales de la entidad financiera.",
    ]:
        add_bullet(doc, item)

    path = os.path.join(BASE, "Software Design Document (SDD).docx")
    doc.save(path)
    print(f"SDD guardado: {path}")


if __name__ == "__main__":
    build_srs()
    build_sdd()
    print("Documentos generados correctamente.")
